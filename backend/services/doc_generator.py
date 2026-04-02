import io
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Pt

def generate_pdf(content: str, header_info: dict) -> bytes:
    """
    Generates a high-fidelity PDF cover letter.
    header_info contains: name, email, phone, location, date.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=LETTER, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    
    # Custom Styles
    name_style = ParagraphStyle(
        'HeaderName',
        parent=styles['Heading1'],
        fontSize=24,
        leading=28,
        alignment=TA_LEFT,
        spaceAfter=4,
        textColor='#2D3748'
    )
    
    contact_style = ParagraphStyle(
        'ContactInfo',
        parent=styles['Normal'],
        fontSize=10,
        leading=12,
        alignment=TA_LEFT,
        textColor='#718096',
        spaceAfter=2
    )

    body_style = ParagraphStyle(
        'BodyText',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        alignment=TA_LEFT,
        spaceBefore=12,
        spaceAfter=12
    )

    elements = []
    
    # Header
    elements.append(Paragraph(header_info.get("name", "Applicant Name"), name_style))
    elements.append(Paragraph(f"{header_info.get('location', '')} | {header_info.get('phone', '')} | {header_info.get('email', '')}", contact_style))
    elements.append(Spacer(1, 24))
    elements.append(Paragraph(header_info.get("date", ""), contact_style))
    elements.append(Spacer(1, 24))
    
    # Body
    # Split content by double newlines for paragraphs
    paragraphs = content.split('\n\n')
    for p_text in paragraphs:
        if p_text.strip():
            elements.append(Paragraph(p_text.replace('\n', '<br/>'), body_style))
    
    doc.build(elements)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes

def generate_docx(content: str, header_info: dict) -> bytes:
    """Generates a standard Word document (.docx) cover letter."""
    doc = Document()
    
    # Header
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(header_info.get("name", "Applicant Name"))
    name_run.bold = True
    name_run.font.size = Pt(18)
    
    contact = f"{header_info.get('location', '')} | {header_info.get('phone', '')} | {header_info.get('email', '')}"
    contact_para = doc.add_paragraph(contact)
    contact_para.style.font.size = Pt(10)
    
    doc.add_paragraph(header_info.get("date", ""))
    doc.add_paragraph() # Spacer
    
    # Body
    paragraphs = content.split('\n\n')
    for p_text in paragraphs:
        if p_text.strip():
            p = doc.add_paragraph(p_text.strip())
            p.style.font.size = Pt(11)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    return docx_bytes
